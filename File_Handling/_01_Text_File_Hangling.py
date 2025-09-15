"text File handling"

"""
open()      # To open a file
read()      # To read file contents
write()     # To write to a file
close()     # To close the file

MODE   | EXPLANATION                                                      | EXAMPLE
-------|------------------------------------------------------------------|---------------------------
r      | Read mode (default). File must exist.                            | open("file.txt", "r")
w      | Write mode. Overwrites if file exists or creates new file.       | open("file.txt", "w")
x      | Exclusive creation. Fails if file exists.                        | open("file.txt", "x")
a      | Append mode. Adds to file or creates if not present.             | open("file.txt", "a")
r+     | Read & write. File must exist. Doesn't erase content.            | open("file.txt", "r+")
w+     | Write & read. Overwrites existing file or creates new.           | open("file.txt", "w+")
a+     | Append & read. Writes at end, reads from start.                  | open("file.txt", "a+")
x+     | Create new file for read & write. Fails if file exists.          | open("file.txt", "x+")

"""

# 'x' = Create a new file, fail if it exists
# file = open('E:\Data Engineer\PythonGit\File_Handling\example.txt', 'x')
# content = file.write("Hello, world!")

# 'r' = Read-only mode
# file = open("example.txt", "r")
# content = file.read()
# print(content)
# file.close()

# 'w' = Write to a File (Overwrites)
# file = open("example.txt", "w")
# file.write("Hello, file handling in Python!")
# file.close()


# 'a' = Append to a File
# file = open("example.txt", "a")
# file.write("\nAdding more content.")
# file.close()


"with is a context manager that automatically manages resources.it automatically opens and closes the file for you."
# with open("example.txt", "r") as file:
#     data = file.read()
#     print(data)


# 'r+' — Read and write (file must exist)
# with open("example.txt", "r+") as file:
#     content = file.read(5)  # Read first 5 characters
#     print("Read:", content)  # Output: Hello
#
#     file.write(" Python")  # Write " Python" starting at current position (after "Hello")
#
#     file.seek(0)  # Go back to start
#     print("After write:", file.read())  # Read entire updated content


# 'w+' — Write and read (creates or truncates file)

# with open("example2.txt", "w+") as file:
#     file.write("Hello World!\n")  # Write text to the empty file
#
#     file.seek(0)  # Move pointer to beginning before reading
#     content = file.read()  # Read entire content
#     print("Content:", content)

"~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~"
"pdf File Handling"

import fitz  # PyMuPDF

# doc = fitz.open()
# page = doc.new_page()
# page.insert_text((72, 72), "Hello, this is a new PDF!")
# doc.save("created.pdf")

# Read a PDF
# doc = fitz.open("E:\Data Engineer\PythonGit\File_Handling\created.pdf")
# text = doc[0].get_text()
# print("Text on page 1:", text)
# doc.close()

# Update a PDF (Add Text)
# doc = fitz.open("created.pdf")
# page = doc[0]
# page.insert_text((72, 100), "This text was added later!")
# doc.save("updated.pdf")

# Delete a Page
# doc = fitz.open("updated.pdf")
# doc.delete_page(0)  # Delete the first page (0-indexed)
# doc.save("after_delete.pdf")

"""
# Write a Python program to read a text file and count the frequency of each word.
# Handle cases where the file might not exist.
# """

# try:
#     with open("mock.txt","r") as f:
#         res =  f.read()
#         str_ls = res.strip().split()
#         dc_ls ={}
#         for ele in str_ls :
#             if ele not  in dc_ls:
#                 dc_ls[ele] = 1
#             else:
#                 dc_ls[ele] += 1
#         print(dc_ls)
# except FileNotFoundError as fe:
#     print(fe)

"""
Create a function that tries to open a file,
and if it doesn't exist, it should create the file with some default content.
"""

# def open_file_fun(f,c):
#     try:
#         with open(f,"r") as f:
#             res = f.read()
#             print(res)
#     except FileNotFoundError as fe:
#         with open(f,"w+") as f:
#             f.write(c)
#             f.seek(0)
#             res = f.read()
#             print(res)
#         print(fe)
#
#
# file_name = "mock.txt"
# content = "Hello Python"
# open_file_fun(file_name,content)


"""
 Write a program that:
   - Opens a file
   - Reads integers line by line
   - Tries to divide 100 by each integer
   - Catches ZeroDivisionError, ValueError, and FileNotFoundError.
"""


# def divide_file_element():
#     try:
#         with open("divide_file_element.txt", "r") as f:
#             main_text = f.read()
#             ls_ele = main_text.strip().split()
#             for ele in ls_ele:
#                 try:
#                     res = 100/int(ele)
#                     print(res)
#                 except (ZeroDivisionError, ValueError) as er:
#                     print(er)
#     except FileNotFoundError as ffe:
#         print("FileNotFoundError :",ffe)
#
#
# divide_file_element()


"~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~"
#CSV file handling (Comma-Separated Values)

#writer() writes rows to a CSV file. Takes a list as input for each row
# import csv
#
# with open("data.csv", mode='w', newline='') as file:
#     writer = csv.writer(file)
#     writer.writerow(["id", "name", "age"])
#     writer.writerow([1, "Alice", 30])
#     writer.writerow([2, "Bob", 25])


# reader() reads CSV files row by row. Each row is returned as a list.

# with open("E:\Data Engineer\PythonGit\File_Handling\data.csv",mode= 'r') as file:
#     reader = csv.reader(file)
#     for row in reader:
#         print(row)

# pandas.read_csv() loads a CSV into a DataFrame, allowing advanced manipulation.
# import pandas as pd
# df = pd.read_csv("E:\Data Engineer\PythonGit\File_Handling\data.csv")
# print(df)
# df.to_csv("new1data.csv", index= False)


"~~~~~~~~~~~~~~~~~~~~~~~~~ CRUD OPERATIONS IN CSV FILES using CSV ~~~~~~~~~~~~~~~~~~~~~~~~~"

# CREATE (Add data)
# Using csv:
# import csv
# with open("csvdata.csv", "w", newline='') as file:
#     writer = csv.writer(file)
#     writer.writerow(["id", "name", "age"])
#     writer.writerow([1, "Alice", 30])
#     writer.writerow([2, "Bob", 25])

# READ (Load and display data)
# Using csv:
# import csv
# with open("csvdata.csv", "r") as file:
#     reader = csv.reader(file)
#     for row in reader:
#         print(row)

# UPDATE (Modify existing row)
# Using csv:
# import csv
# rows = []
# with open("data.csv", "r") as file:
#     reader = csv.reader(file)
#     for row in reader:
#         if row[0] == "1":
#             row[2] = "31"
#         rows.append(row)
# with open("data.csv", "w", newline='') as file:
#     writer = csv.writer(file)
#     writer.writerows(rows)



# DELETE (Remove a row)
# Using csv:
# import csv
# rows = []
# with open("csvdata.csv", "r") as file:
#     reader = csv.reader(file)
#     for row in reader:
#         if row[0] != "2":
#             rows.append(row)
# with open("csvdata.csv", "w", newline='') as file:
#     writer = csv.writer(file)
#     writer.writerows(rows)


"~~~~~~~~~~~~~~~~~~~~~~~~~ CRUD OPERATIONS IN CSV FILES using Panda ~~~~~~~~~~~~~~~~~~~~~~~~~"

# CREATE (Add data)
# Using pandas:
# import pandas as pd
# data = {"id": [1, 2], "name": ["Alice", "Bob"], "age": [30, 25]}
# df = pd.DataFrame(data)
# df.to_csv("pandadata.csv", index=False)

# READ (Load and display data)
# Using pandas:
# import pandas as pd
# df = pd.read_csv("pandadata.csv")
# print(df)

# UPDATE (Modify existing row)
# Using pandas:
# import pandas as pd
# df = pd.read_csv("E:\Data Engineer\PythonGit\File_Handling\pandadata.csv")
# df.loc[df['id'] == 1, 'age'] = 31
# df.to_csv("E:\Data Engineer\PythonGit\File_Handling\pandadata.csv", index=False)

# DELETE (Remove a row)
# Using pandas:
# import pandas as pd
# df = pd.read_csv("E:\Data Engineer\PythonGit\File_Handling\pandadata.csv")
# df = df[df['id'] != 2]
# df.to_csv("E:\Data Engineer\PythonGit\File_Handling\pandadata.csv", index=False)


"================ JSON FILE HANDLING ========================================="
# with open("E:\Data Engineer\PythonGit\File_Handling\jsondata.json", "r") as file:
#     data = file.read()
#     print(data)



# json.load() reads and parses a JSON file into a Python dictionary or list. Example:
# import json
# with open("E:\Data Engineer\PythonGit\File_Handling\jsondata.json", "r") as file:
#     data = json.load(file)
#     print(data)


# json.dump() writes a Python dictionary/list into a JSON file.
# import json
#
# # Step 1: New data to add
# new_user = {
#     "id": 4,
#     "name": "David",
#     "age": 28,
#     "email": "david@example.com"
# }
#
# # Step 2: Load existing data from file
# with open("E:\Data Engineer\PythonGit\File_Handling\jsondata.json", "r") as file:
#     existing_data = json.load(file)
#
# # Step 3: Append new user
# existing_data.append(new_user)
#
# # Step 4: Write updated data back to file
# with open("E:\Data Engineer\PythonGit\File_Handling\jsondata.json", "w") as file:
#     json.dump(existing_data, file, indent=4)


# json.loads() converts a JSON string into a Python dictionary.
# import json
# json_string = '{"name": "Alice", "age": 30}'
# data = json.loads(json_string)
# print(data,type(data))

# json.dumps() converts a Python dictionary into a JSON-formatted string
# import json
# data = {"name": "Alice", "age": 30}
# json_string = json.dumps(data)
# print(data,type(json_string))


"============= CRUD OPERATIONS IN JSON FILES ======================"

# CREATE (Write new JSON data to a file)
# import json
# data = {"id": 1, "name": "Alice", "age": 30}
# with open("json_crud_data.json", "w") as file:
#     json.dump(data, file, indent=4)


# READ (Read JSON data from a file)
# import json
# with open("E:\Data Engineer\PythonGit\File_Handling\json_crud_data.json", "r") as file:
#     data = json.load(file)
#     print(data)


# # UPDATE (Modify values in existing JSON data)
# import json
# with open("E:\Data Engineer\PythonGit\File_Handling\json_crud_data.json", "r") as file:
#     data = json.load(file)
# data["age"] = 18
# with open("E:\Data Engineer\PythonGit\File_Handling\json_crud_data.json", "w") as file:
#     json.dump(data, file, indent=4)
#
#
# # DELETE (Remove key or entry from JSON)
# import json
# with open("json_crud_data.json", "r") as file:
#     data = json.load(file)
# del data["age"]
# with open("json_crud_data.json", "w") as file:
#     json.dump(data, file, indent=4)

"~~~~~~~~~~~~~~~~~~~ CRUD OPERATIONS IN JSON FILES using Pandas ~~~~~~~~~~~~~~~~~~~~~~~~~~~~"
# CREATE (Add new record to JSON)
# import pandas as pd
# df = pd.read_json("pandas_json_crud_data.json")
#
# new_record = pd.DataFrame([{"id": 1, "name": "Alice", "age": 30}])
# df = pd.concat([df, new_record], ignore_index=True)
# df.to_json("pandas_json_crud_data.json", orient="records", indent=4)


# READ (Load JSON data into DataFrame)
# import pandas as pd
# df = pd.read_json("pandas_json_crud_data.json")
# print(df)

# UPDATE (Modify values in JSON data)

# import pandas as pd
#
# df = pd.read_json("pandas_json_crud_data.json")
# df.loc[df['id'] == 1, 'age'] = 55
# df.to_json("pandas_json_crud_data.json", orient="records", indent=4)

# DELETE (Remove record or column from JSON)
# import pandas as pd
# df = pd.read_json("pandas_json_crud_data.json")
# # df = df[df['id'] != 1]
# # Or delete a column (e.g., 'age')
# df = df.drop(columns=['age'])
# df.to_json("pandas_json_crud_data.json", orient="records", indent=4)


"============================= DOCX FILE HANDLING ======================"
from docx import Document
# Document() is the main class for creating or loading Word documents.
# To create a new document:
# doc = Document("file.docx")


# open() is not used in the traditional sense for .docx. Instead, `docx.Document()` is used to load or create documents.
# Example (to open):
# doc = Document("file.docx")
# print(doc)


# read() is the concept of accessing content, such as paragraphs, headings, and tables. You read them using iteration.
# Example:
# for para in doc.paragraphs:
#     print(para.text)

# write() refers to adding text or elements like headings, tables, and paragraphs to a document.
# Example:
# doc.add_paragraph("This is new content.")


# save() is used to save changes made to the document to a file.
# Example:
# doc.save("updated.docx")


# Document() is the main class for creating or loading Word documents.
# To create a new document:
# doc = Document()


# add_paragraph() adds a new paragraph to the document.
# Example:
# doc.add_paragraph("Hello, World!")

# add_heading() adds a heading of level 0 to 4.
# Example:
# doc.add_heading("Document Title", level=2)


"============================ CRUD OPERATIONS IN .DOCX FILES =================================="
from docx import Document
# CREATE (Create and write to a new document)
#
# doc = Document()
# doc.add_heading("Employee Report", level=1)
# doc.add_paragraph("Name: Alice")
# doc.add_paragraph("Department: HR")
# doc.save("docx_crud_report.docx")


# # READ (Open and extract data from an existing document)
# d = Document("docx_crud_report.docx")
# for para in d.paragraphs:
#     print(para.text)

# UPDATE (Modify existing content or add new content)
# doc = Document("docx_crud_report.docx")
# doc.add_paragraph("Status: Active")
# doc.paragraphs[1].text = "Name: Alice Smith"
# doc.save("docx_crud_report.docx")

# DELETE (No direct delete API, but can be mimicked by recreating content)
# doc = Document("docx_crud_report.docx")
# new_doc = Document()
# for i, para in enumerate(doc.paragraphs):
#     if "Status" not in para.text:
#         new_doc.add_paragraph(para.text)
# new_doc.save("docx_crud_report.docx")



# DOCX TABLE EXAMPLE (Create a document with a table)
# doc = Document()
# table = doc.add_table(rows=3, cols=3)
# table.style = 'Table Grid'
# table.cell(0, 0).text = "ID"
# table.cell(0, 1).text = "Name"
# table.cell(0, 2).text = "Dept"
# table.cell(1, 0).text = "1"
# table.cell(1, 1).text = "Alice"
# table.cell(1, 2).text = "HR"
# doc.save("docx_crud_report.docx")
#
#
# "==================== XLSX FILE HANDLING IN PYTHON ============================"
# "==================== CRUD XLSX FILE HANDLING  from openpyxl import Workbook ====================="
# from openpyxl import Workbook
# wb = Workbook()
# ws = wb.active
# ws.title = "Employees"
# ws.append(["ID", "Name", "Dept"])
# ws.append([1, "Alice", "HR"])
# ws.append([2, "Bob", "IT"])
# wb.save("employees.xlsx")



